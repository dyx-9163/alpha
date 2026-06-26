from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path("outputs/algeria_openlab_inventory.docx")


HEADERS = [
    "虚拟机名",
    "系统",
    "厂家",
    "版本",
    "资源规格",
    "系统",
    "弹性IP地址",
    "内网IP地址",
    "系统后台登陆账号/密码",
    "系统访问信息",
]


ROWS = [
    [
        "致远伙伴应用IM_415646",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：500GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.53",
        "10.248.1.240",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用IM_135988",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：500GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.35",
        "10.248.1.160",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用IM_231975",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：500GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.30",
        "10.248.1.92",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用IM_265680",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：500GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.38",
        "10.248.1.79",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用TQN_777497",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：4核\nmem：16GB\nDisk：200GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.25",
        "10.248.1.120",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用\nDataBase_856988",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：200GB*3+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.28",
        "10.248.1.234",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用\nDataBase_876747",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：200GB*3+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.57",
        "10.248.1.216",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用app02_521599",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：200GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.9",
        "10.248.1.75",
        "root/Image0@Huawei123",
        "",
    ],
    [
        "致远伙伴应用app01_873529",
        "HCS_851",
        "北京致远互联软件股份有限公司",
        "openEuler-24.03-LTS-SP3-x86_64",
        "vCPU：16核\nmem：32GB\nDisk：200GB+100GB",
        "openEuler-24.03-LTS-SP3-x86_64-电力专用",
        "172.21.248.21",
        "10.248.1.214",
        "root/Image0@Huawei123",
        "",
    ],
]


COL_WIDTHS = [1.72, 0.55, 0.88, 1.02, 1.40, 1.12, 0.95, 1.02, 1.35, 1.20]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="7A8793", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size_pt=8.0, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def fill_cell(cell, text, size_pt=8.0, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    parts = str(text).split("\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(part)
        set_font(run, size_pt=size_pt, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_cell_width(cell, width_in):
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    tc_w.type = "dxa"
    tc_w.w = int(width_in * 1440)


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("阿尔及利亚环境交接及云上 Openlab 使用指导")
    set_font(run, size_pt=16, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("服务器环境清单（根据截图整理）")
    set_font(run, size_pt=10, color="555555")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run("说明：以下内容由截图转录为可编辑表格；系统访问信息列截图中未填写，已保留空白。账号/密码等敏感字段请在交付前复核并按安全规范保存。")
    set_font(run, size_pt=9, color="555555")

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(HEADERS):
        cell = hdr.cells[i]
        set_cell_width(cell, COL_WIDTHS[i])
        set_cell_shading(cell, "D9EAF7")
        fill_cell(cell, header, size_pt=8.0, bold=True, color="1F2937")

    for row_index, row_data in enumerate(ROWS, start=1):
        row = table.add_row()
        for col_index, value in enumerate(row_data):
            cell = row.cells[col_index]
            set_cell_width(cell, COL_WIDTHS[col_index])
            set_cell_shading(cell, "EEF6FC" if row_index % 2 else "FFFFFF")
            align = WD_ALIGN_PARAGRAPH.LEFT if col_index in (0, 2, 3, 4, 5, 8, 9) else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cell, value, size_pt=7.2, align=align)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(8)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("来源：会议截图，截图时间 2026/6/15 11:11")
    set_font(run, size_pt=8.5, color="666666")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
